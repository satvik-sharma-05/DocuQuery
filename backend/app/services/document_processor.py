import io
import PyPDF2
from docx import Document
from typing import List, BinaryIO
import logging
import re
from app.core.config import settings

logger = logging.getLogger(__name__)

# Safety limits
MAX_CHUNKS = 2000
MAX_TEXT_LENGTH = 10_000_000  # 10MB of text

class DocumentProcessor:
    def __init__(self):
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
        
        # Validate configuration
        if self.chunk_overlap >= self.chunk_size:
            logger.warning(
                f"chunk_overlap ({self.chunk_overlap}) >= chunk_size ({self.chunk_size}). "
                f"Setting overlap to chunk_size - 100"
            )
            self.chunk_overlap = max(0, self.chunk_size - 100)
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file with enhanced formatting preservation"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_parts = []
            
            # Extract paragraphs with better formatting
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    # Preserve paragraph structure
                    text_parts.append(para_text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            # Join with double newlines to preserve document structure
            full_text = "\n\n".join(text_parts)
            
            # Clean up excessive whitespace while preserving structure
            full_text = re.sub(r'\n\s*\n\s*\n+', '\n\n', full_text)
            full_text = re.sub(r'[ \t]+', ' ', full_text)
            
            return full_text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise
    
    def extract_text_from_txt(self, file_content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            return file_content.decode('utf-8').strip()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                return file_content.decode('latin-1').strip()
            except Exception as e:
                logger.error(f"Error extracting text from TXT: {str(e)}")
                raise
    
    def extract_text(self, file_content: bytes, file_extension: str) -> str:
        """Extract text based on file extension"""
        file_extension = file_extension.lower().lstrip('.')
        
        if file_extension == 'pdf':
            return self.extract_text_from_pdf(file_content)
        elif file_extension == 'docx':
            return self.extract_text_from_docx(file_content)
        elif file_extension in ['txt', 'md']:
            return self.extract_text_from_txt(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def chunk_text(self, text: str) -> List[str]:
        """
        Intelligently split text into overlapping chunks while preserving context.
        """
        # Validate input
        if not text or not text.strip():
            return []
        
        text_length = len(text)
        
        # Safety check for text length
        if text_length > MAX_TEXT_LENGTH:
            logger.warning(f"Text too long ({text_length} chars), truncating to {MAX_TEXT_LENGTH}")
            text = text[:MAX_TEXT_LENGTH]
            text_length = MAX_TEXT_LENGTH
        
        # Log chunking parameters
        logger.info(f"Text length: {text_length} characters")
        logger.info(f"Chunk size: {self.chunk_size}")
        logger.info(f"Overlap: {self.chunk_overlap}")
        
        # Validate overlap
        if self.chunk_overlap >= self.chunk_size:
            raise ValueError(
                f"overlap ({self.chunk_overlap}) must be smaller than chunk_size ({self.chunk_size})"
            )
        
        # If text is smaller than chunk size, return as single chunk
        if text_length <= self.chunk_size:
            return [text.strip()]
        
        chunks = []
        start = 0
        step = self.chunk_size - self.chunk_overlap
        
        # Ensure step is positive to prevent infinite loop
        if step <= 0:
            step = self.chunk_size // 2
            logger.warning(f"Invalid step size, using {step}")
        
        while start < text_length:
            end = min(start + self.chunk_size, text_length)
            
            # Smart boundary detection - try multiple strategies
            if end < text_length:
                # Strategy 1: Look for paragraph breaks (double newlines)
                paragraph_boundary = text.rfind('\n\n', start, end)
                if paragraph_boundary > start:
                    end = paragraph_boundary + 2
                else:
                    # Strategy 2: Look for sentence boundaries
                    sentence_boundaries = ['. ', '! ', '? ', '.\n', '!\n', '?\n']
                    best_boundary = start
                    
                    for punct in sentence_boundaries:
                        boundary = text.rfind(punct, start, end)
                        if boundary > best_boundary:
                            best_boundary = boundary + len(punct)
                    
                    if best_boundary > start:
                        end = best_boundary
                    else:
                        # Strategy 3: Look for line breaks
                        line_boundary = text.rfind('\n', start, end)
                        if line_boundary > start:
                            end = line_boundary + 1
                        else:
                            # Strategy 4: Look for word boundaries
                            word_boundary = text.rfind(' ', start, end)
                            if word_boundary > start:
                                end = word_boundary
            
            chunk = text[start:end].strip()
            
            if chunk:
                # Ensure chunk has meaningful content (not just whitespace/punctuation)
                if len(chunk.replace(' ', '').replace('\n', '').replace('\t', '')) > 10:
                    chunks.append(chunk)
            
            # Safety check: prevent infinite loop
            if len(chunks) > MAX_CHUNKS:
                logger.warning(f"Reached maximum chunk limit ({MAX_CHUNKS}), truncating")
                break
            
            # Move forward by step size (CRITICAL: must always move forward)
            start += step
            
            # Additional safety: ensure we're making progress
            if start <= (end - step):
                logger.error(f"Chunking not making progress, breaking loop")
                break
        
        logger.info(f"Created {len(chunks)} chunks")
        
        return chunks

# Global document processor instance
document_processor = DocumentProcessor()